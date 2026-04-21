import json
from datetime import datetime, date, timedelta
import io
import re

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select, func, and_, cast, Date, case
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

from app.core.database import get_db
from app.core.response import ApiResponse
from app.core.security import decode_access_token, oauth2_scheme
from app.models.models import Task, TaskCategory, DailyCheckin, TaskPlanTemplate, TaskPlanEntry
from app.schemas.schemas import (
    TaskCreate, TaskUpdate, TaskOut,
    TaskCategoryCreate, TaskCategoryUpdate, TaskCategoryOut,
    DailyCheckinCreate, DailyCheckinOut,
    TaskPlanTemplateOut,
)


class PlanEntryUpdatePayload(BaseModel):
    planned_text: str | None = None
    manual_text: str | None = None


class PlanExportMappingPayload(BaseModel):
    date_col: str = "日期"
    plan_col: str = "计划内容"
    auto_col: str = "自动回填"
    manual_col: str = "手工补充"
    final_col: str = "最终上交内容"
    rate_col: str = "完成率"
    lock_col: str = "锁定"

router = APIRouter()


async def get_current_user_id(token: str = Depends(oauth2_scheme)) -> str:
    payload = decode_access_token(token)
    user_id = payload.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="无效的认证")
    return user_id


def _normalize_cell_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _parse_day(value: str) -> int | None:
    m = re.search(r"(\d{1,2})", value or "")
    if not m:
        return None
    d = int(m.group(1))
    if 1 <= d <= 31:
        return d
    return None


def _extract_plan_rows_from_docx(content: bytes) -> list[dict]:
    doc = Document(io.BytesIO(content))
    rows: list[dict] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize_cell_text(c.text) for c in row.cells]
            if not cells:
                continue
            day = _parse_day(cells[0])
            if day is None:
                continue
            planned_text = ""
            if len(cells) >= 2:
                planned_text = cells[1]
            elif len(cells) == 1:
                planned_text = ""
            rows.append({"day": day, "planned_text": planned_text})
    dedup = {}
    for item in rows:
        dedup[item["day"]] = item["planned_text"]
    return [{"day": d, "planned_text": dedup[d]} for d in sorted(dedup.keys())]


def _infer_export_mapping_from_docx(content: bytes) -> dict:
    mapping = {
        "date_col": "日期",
        "plan_col": "计划内容",
        "auto_col": "自动回填",
        "manual_col": "手工补充",
        "final_col": "最终上交内容",
        "rate_col": "完成率",
        "lock_col": "锁定",
    }
    doc = Document(io.BytesIO(content))
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [_normalize_cell_text(c.text) for c in table.rows[0].cells]
        for cell in header_cells:
            if not cell:
                continue
            if any(k in cell for k in ("日期", "时间", "day")):
                mapping["date_col"] = cell
            elif any(k in cell for k in ("计划", "安排", "任务")):
                mapping["plan_col"] = cell
            elif any(k in cell for k in ("完成率", "进度")):
                mapping["rate_col"] = cell
            elif any(k in cell for k in ("锁定", "冻结")):
                mapping["lock_col"] = cell
            elif any(k in cell for k in ("实际", "回填", "达成")):
                mapping["final_col"] = cell
            elif any(k in cell for k in ("上午", "早")):
                mapping["morning_col"] = cell
            elif any(k in cell for k in ("下午", "午后")):
                mapping["afternoon_col"] = cell
            elif any(k in cell for k in ("晚间", "晚上", "夜间")):
                mapping["evening_col"] = cell
    return mapping


def _extract_slot_text(actual_text: str | None, slot: str) -> str:
    if not actual_text:
        return ""
    for line in actual_text.splitlines():
        if line.startswith(f"{slot} 已完成："):
            return line
    return ""


def _build_plan_excel(template: TaskPlanTemplate, entries: list[TaskPlanEntry]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "月计划"
    mapping = {
        "date_col": "日期",
        "plan_col": "计划内容",
        "auto_col": "自动回填",
        "manual_col": "手工补充",
        "final_col": "最终上交内容",
        "rate_col": "完成率",
        "lock_col": "锁定",
    }
    if template.export_mapping:
        try:
            parsed = json.loads(template.export_mapping)
            if isinstance(parsed, dict):
                mapping.update({k: str(v) for k, v in parsed.items() if k in mapping})
        except Exception:
            pass
    ws.append([f"{template.month} 备考计划回填表"])
    ws.append([
        mapping["date_col"], mapping["plan_col"], mapping["auto_col"], mapping["manual_col"],
        mapping["final_col"], mapping["rate_col"], mapping["lock_col"],
    ])
    for e in sorted(entries, key=lambda x: x.day):
        date_text = f"{template.month}-{str(e.day).zfill(2)}"
        final_text = "；".join([txt for txt in [e.actual_text or "", e.manual_text or ""] if txt]).strip("；")
        ws.append([
            date_text,
            e.planned_text or "",
            e.actual_text or "",
            e.manual_text or "",
            final_text,
            f"{e.completion_rate}%" if e.completion_rate is not None else "",
            "是" if e.locked else "否",
        ])
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 26
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 26
    ws.column_dimensions["E"].width = 36
    ws.column_dimensions["F"].width = 10
    ws.column_dimensions["G"].width = 8
    title_font = Font(name="Microsoft YaHei", bold=True, size=14)
    header_font = Font(name="Microsoft YaHei", bold=True, size=10)
    body_font = Font(name="Microsoft YaHei", size=10)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_wrap = Alignment(horizontal="left", vertical="top", wrap_text=True)
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill(fill_type="solid", start_color="EEF2FF", end_color="EEF2FF")

    ws["A1"].font = title_font
    ws["A1"].alignment = center
    for col in range(1, 8):
        header_cell = ws.cell(row=2, column=col)
        header_cell.font = header_font
        header_cell.fill = header_fill
        header_cell.alignment = center
        header_cell.border = border
    for row in range(3, ws.max_row + 1):
        for col in range(1, 8):
            c = ws.cell(row=row, column=col)
            c.font = body_font
            c.border = border
            c.alignment = center if col in (1, 6, 7) else left_wrap
        ws.row_dimensions[row].height = 40
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def _summarize_tasks_for_day(tasks: list[Task]) -> tuple[str, float]:
    total = len(tasks)
    completed_tasks = [t for t in tasks if t.status == "completed"]
    completion_rate = round((len(completed_tasks) / total) * 100, 1) if total > 0 else 0.0

    def _slot(dt: datetime | None) -> str:
        if not dt:
            return "未设时段"
        h = dt.hour
        if h < 12:
            return "上午"
        if h < 18:
            return "下午"
        return "晚间"

    slots: dict[str, dict[str, list[str]]] = {
        "上午": {"done": [], "todo": []},
        "下午": {"done": [], "todo": []},
        "晚间": {"done": [], "todo": []},
        "未设时段": {"done": [], "todo": []},
    }
    for t in tasks:
        s = _slot(t.due_date)
        bucket = "done" if t.status == "completed" else "todo"
        slots[s][bucket].append(t.title)

    lines = [f"完成 {len(completed_tasks)}/{total} 项"]
    for name in ("上午", "下午", "晚间", "未设时段"):
        done = "、".join(slots[name]["done"]) if slots[name]["done"] else "无"
        todo = "、".join(slots[name]["todo"]) if slots[name]["todo"] else "无"
        lines.append(f"{name} 已完成：{done}；未完成：{todo}")
    return "\n".join(lines), completion_rate


def _template_payload(template: TaskPlanTemplate, entries: list[TaskPlanEntry]) -> dict:
    return TaskPlanTemplateOut(
        id=template.id,
        name=template.name,
        month=template.month,
        source_filename=template.source_filename,
        export_mapping=template.export_mapping,
        entries=entries,
        created_at=template.created_at,
        updated_at=template.updated_at,
    ).model_dump()


# ═══════════════════════════════════════
# Task Categories
# ═══════════════════════════════════════

@router.post("/categories", status_code=201)
async def create_category(
    body: TaskCategoryCreate,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    cat = TaskCategory(user_id=user_id, **body.model_dump())
    db.add(cat)
    await db.flush()
    await db.refresh(cat)
    return ApiResponse.success(data=TaskCategoryOut.model_validate(cat).model_dump())


@router.get("/categories")
async def list_categories(
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskCategory).where(TaskCategory.user_id == user_id)
            .order_by(TaskCategory.sort_order, TaskCategory.created_at)
    )
    cats = result.scalars().all()
    return ApiResponse.success(data=[TaskCategoryOut.model_validate(c).model_dump() for c in cats])


@router.put("/categories/{cat_id}")
async def update_category(
    cat_id: str,
    body: TaskCategoryUpdate,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskCategory).where(and_(TaskCategory.id == cat_id, TaskCategory.user_id == user_id))
    )
    cat = result.scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="分类不存在")
    for k, v in body.model_dump(exclude_unset=True).items():
        setattr(cat, k, v)
    await db.flush()
    await db.refresh(cat)
    return ApiResponse.success(data=TaskCategoryOut.model_validate(cat).model_dump())


@router.delete("/categories/{cat_id}")
async def delete_category(
    cat_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskCategory).where(and_(TaskCategory.id == cat_id, TaskCategory.user_id == user_id))
    )
    cat = result.scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="分类不存在")
    # Reset tasks in this category to "其他"
    tasks = (await db.execute(
        select(Task).where(and_(Task.user_id == user_id, Task.category == cat.name))
    )).scalars().all()
    for t in tasks:
        t.category = "其他"
    await db.delete(cat)
    return ApiResponse.success(message="已删除")


# ═══════════════════════════════════════
# Tasks CRUD
# ═══════════════════════════════════════

@router.post("/", status_code=201)
async def create_task(
    body: TaskCreate,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    task = Task(
        user_id=user_id,
        title=body.title,
        description=body.description,
        priority=body.priority,
        status=body.status,
        category=body.category,
        due_date=body.due_date,
        estimated_time=body.estimated_time,
        tags=json.dumps(body.tags) if body.tags else None,
        subtasks=json.dumps([s.model_dump() for s in body.subtasks]) if body.subtasks else None,
        is_pinned=body.is_pinned,
    )
    db.add(task)
    await db.flush()
    await db.refresh(task)
    return ApiResponse.success(data=TaskOut.model_validate(task).model_dump())


@router.get("/")
async def list_tasks(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=500),
    status: str | None = Query(None, pattern="^(pending|in_progress|completed)$"),
    category: str | None = None,
    priority: str | None = Query(None, pattern="^(high|medium|low)$"),
    search: str | None = None,
    sort_by: str = Query("created_at", pattern="^(created_at|due_date|priority|updated_at)$"),
    sort_order: str = Query("desc", pattern="^(asc|desc)$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    query = select(Task).where(Task.user_id == user_id)
    count_query = select(func.count()).select_from(Task).where(Task.user_id == user_id)

    if status:
        query = query.where(Task.status == status)
        count_query = count_query.where(Task.status == status)
    if category:
        query = query.where(Task.category == category)
        count_query = count_query.where(Task.category == category)
    if priority:
        query = query.where(Task.priority == priority)
        count_query = count_query.where(Task.priority == priority)
    if search:
        search_filter = Task.title.ilike(f"%{search}%")
        query = query.where(search_filter)
        count_query = count_query.where(search_filter)

    # Pinned first, then sort
    if sort_by == "priority":
        col = case(
            (Task.priority == "high", 1),
            (Task.priority == "medium", 2),
            else_=3,
        )
    else:
        col = getattr(Task, sort_by)

    order_col = col.asc() if sort_order == "asc" else col.desc()
    query = query.order_by(Task.is_pinned.desc(), order_col.nullslast())

    total = (await db.execute(count_query)).scalar()
    tasks = (await db.execute(
        query.offset((page - 1) * page_size).limit(page_size)
    )).scalars().all()

    return ApiResponse.success(data={
        "items": [TaskOut.model_validate(t).model_dump() for t in tasks],
        "total": total,
        "page": page,
        "page_size": page_size,
    })


@router.get("/today")
async def get_today_tasks(
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    today = date.today()
    result = await db.execute(
        select(Task).where(
            and_(
                Task.user_id == user_id,
                Task.status != "completed",
            )
        ).order_by(
            Task.is_pinned.desc(),
            case(
                (Task.priority == "high", 1),
                (Task.priority == "medium", 2),
                else_=3,
            ),
            Task.due_date.asc().nullslast(),
        )
    )
    tasks = result.scalars().all()
    # Separate overdue
    today_tasks = []
    overdue_tasks = []
    for t in tasks:
        if t.due_date and t.due_date.date() < today:
            overdue_tasks.append(TaskOut.model_validate(t).model_dump())
        else:
            today_tasks.append(TaskOut.model_validate(t).model_dump())

    return ApiResponse.success(data={"today": today_tasks, "overdue": overdue_tasks})


@router.get("/stats")
async def get_task_stats(
    days: int = Query(30, ge=1, le=365),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    now = datetime.utcnow()
    today = date.today()
    week_start = today - timedelta(days=today.weekday())

    # Overall stats
    total = (await db.execute(
        select(func.count()).select_from(Task).where(Task.user_id == user_id)
    )).scalar()

    completed = (await db.execute(
        select(func.count()).select_from(Task).where(
            and_(Task.user_id == user_id, Task.status == "completed")
        )
    )).scalar()

    pending = (await db.execute(
        select(func.count()).select_from(Task).where(
            and_(Task.user_id == user_id, Task.status == "pending")
        )
    )).scalar()

    in_progress = (await db.execute(
        select(func.count()).select_from(Task).where(
            and_(Task.user_id == user_id, Task.status == "in_progress")
        )
    )).scalar()

    # Category distribution
    cat_result = await db.execute(
        select(Task.category, func.count().label("cnt"))
        .where(Task.user_id == user_id)
        .group_by(Task.category)
        .order_by(func.count().desc())
    )
    categories = [{"name": row[0], "count": row[1]} for row in cat_result.all()]
    category_distribution = {row["name"]: row["count"] for row in categories}

    # Overdue count
    overdue = (await db.execute(
        select(func.count()).select_from(Task).where(
            and_(
                Task.user_id == user_id,
                Task.status != "completed",
                Task.due_date.isnot(None),
                Task.due_date < now,
            )
        )
    )).scalar()

    # Completed counts for today / this week
    today_start = datetime.combine(today, datetime.min.time())
    today_end = datetime.combine(today, datetime.max.time())
    week_start_dt = datetime.combine(week_start, datetime.min.time())
    week_end_dt = datetime.combine(today, datetime.max.time())

    today_completed = (await db.execute(
        select(func.count()).select_from(Task).where(
            and_(
                Task.user_id == user_id,
                Task.status == "completed",
                Task.completed_at.isnot(None),
                Task.completed_at >= today_start,
                Task.completed_at <= today_end,
            )
        )
    )).scalar()

    week_completed = (await db.execute(
        select(func.count()).select_from(Task).where(
            and_(
                Task.user_id == user_id,
                Task.status == "completed",
                Task.completed_at.isnot(None),
                Task.completed_at >= week_start_dt,
                Task.completed_at <= week_end_dt,
            )
        )
    )).scalar()

    # Daily completion for last N days
    since = date.today() - timedelta(days=days)
    daily_result = await db.execute(
        select(
            cast(Task.completed_at, Date).label("day"),
            func.count().label("cnt"),
        ).where(
            and_(
                Task.user_id == user_id,
                Task.status == "completed",
                Task.completed_at.isnot(None),
                Task.completed_at >= datetime.combine(since, datetime.min.time()),
            )
        ).group_by(cast(Task.completed_at, Date))
        .order_by(cast(Task.completed_at, Date))
    )
    daily = {str(row[0]): row[1] for row in daily_result.all()}
    daily_completion = [{"date": k, "count": v} for k, v in daily.items()]

    # Streak calculation
    checkin_result = await db.execute(
        select(DailyCheckin.checkin_date)
        .where(
            and_(
                DailyCheckin.user_id == user_id,
                DailyCheckin.tasks_completed > 0,
            )
        )
        .order_by(DailyCheckin.checkin_date.desc())
        .limit(365)
    )
    dates = sorted([row[0] for row in checkin_result.all()], reverse=True)
    streak = 0
    if dates:
        check_date = date.today()
        for d_str in dates:
            d = date.fromisoformat(d_str)
            if d == check_date:
                streak += 1
                check_date -= timedelta(days=1)
            elif d < check_date:
                break

    return ApiResponse.success(data={
        "total": total,
        "completed": completed,
        "pending": pending,
        "in_progress": in_progress,
        "overdue": overdue,
        "completion_rate": round(completed / total * 100, 1) if total > 0 else 0,
        "categories": categories,
        "category_distribution": category_distribution,
        "daily_completion": daily_completion,
        "today_completed": today_completed,
        "week_completed": week_completed,
        "streak": streak,
    })


# ═══════════════════════════════════════
# Daily Checkin
# ═══════════════════════════════════════

@router.post("/checkin", status_code=201)
async def create_or_update_checkin(
    body: DailyCheckinCreate,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(DailyCheckin).where(
            and_(DailyCheckin.user_id == user_id, DailyCheckin.checkin_date == body.checkin_date)
        )
    )
    checkin = result.scalar_one_or_none()
    if checkin:
        checkin.tasks_completed = body.tasks_completed
        checkin.cards_reviewed = body.cards_reviewed
        checkin.study_minutes = body.study_minutes
        checkin.notes_count = body.notes_count
    else:
        checkin = DailyCheckin(user_id=user_id, **body.model_dump())
        db.add(checkin)
    await db.flush()
    await db.refresh(checkin)
    return ApiResponse.success(data=DailyCheckinOut.model_validate(checkin).model_dump())


@router.get("/checkin")
async def get_checkins(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(DailyCheckin).where(
            and_(
                DailyCheckin.user_id == user_id,
                DailyCheckin.checkin_date.startswith(month),
            )
        ).order_by(DailyCheckin.checkin_date)
    )
    checkins = result.scalars().all()
    return ApiResponse.success(data=[DailyCheckinOut.model_validate(c).model_dump() for c in checkins])


@router.post("/plan-template/import-docx")
async def import_plan_template_docx(
    file: UploadFile = File(...),
    month: str = Form(...),
    name: str = Form("课程计划模板"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 模板")
    if not re.match(r"^\d{4}-\d{2}$", month):
        raise HTTPException(status_code=400, detail="month 必须是 YYYY-MM")

    content = await file.read()
    rows = _extract_plan_rows_from_docx(content)
    inferred_mapping = _infer_export_mapping_from_docx(content)
    if not rows:
        raise HTTPException(status_code=400, detail="未在模板中识别到按天计划行")

    existing = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    old = existing.scalar_one_or_none()
    if old:
        await db.delete(old)
        await db.flush()

    template = TaskPlanTemplate(
        user_id=user_id,
        name=name,
        month=month,
        source_filename=file.filename,
        export_mapping=json.dumps(inferred_mapping, ensure_ascii=False),
    )
    db.add(template)
    await db.flush()
    for item in rows:
        db.add(
            TaskPlanEntry(
                template_id=template.id,
                day=item["day"],
                planned_text=item["planned_text"],
                actual_text=None,
                manual_text=None,
                completion_rate=None,
                locked=False,
            )
        )
    await db.flush()
    await db.refresh(template)
    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    data = _template_payload(template, entries)
    return ApiResponse.success(data=data)


@router.get("/plan-template")
async def get_plan_template(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = result.scalar_one_or_none()
    if not template:
        return ApiResponse.success(data=None)
    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    data = _template_payload(template, entries)
    return ApiResponse.success(data=data)


@router.post("/plan-template/generate-today")
async def generate_today_plan_entry(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    now = datetime.utcnow()
    today = now.date()
    month_prefix = f"{today.year:04d}-{today.month:02d}"
    if month != month_prefix:
        raise HTTPException(status_code=400, detail="只能生成当月当天计划")

    t_result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = t_result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="当月计划模板不存在")

    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id)
    )).scalars().all()
    day_map = {e.day: e for e in entries}
    today_entry = day_map.get(today.day)
    if not today_entry:
        raise HTTPException(status_code=404, detail="模板中不存在今日行")

    for entry in entries:
        if entry.day < today.day and entry.actual_text:
            entry.locked = True

    if today_entry.locked and today_entry.actual_text:
        # 保持不变，避免历史内容被覆盖
        pass
    else:
        day_start = datetime.combine(today, datetime.min.time())
        day_end = datetime.combine(today, datetime.max.time())
        tasks = (await db.execute(
            select(Task).where(
                and_(
                    Task.user_id == user_id,
                    Task.due_date.is_not(None),
                    Task.due_date >= day_start,
                    Task.due_date <= day_end,
                )
            )
        )).scalars().all()
        actual_text, completion_rate = _summarize_tasks_for_day(tasks)
        today_entry.actual_text = actual_text
        today_entry.completion_rate = completion_rate
        today_entry.locked = True

    await db.flush()
    refreshed = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    data = _template_payload(template, refreshed)
    return ApiResponse.success(data=data)


@router.post("/plan-template/generate-month")
async def generate_month_plan_entries(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    t_result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = t_result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="当月计划模板不存在")

    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    year, mon = map(int, month.split("-"))
    for entry in entries:
        if entry.locked and entry.actual_text:
            continue
        entry_day = date(year, mon, entry.day)
        day_start = datetime.combine(entry_day, datetime.min.time())
        day_end = datetime.combine(entry_day, datetime.max.time())
        tasks = (await db.execute(
            select(Task).where(
                and_(
                    Task.user_id == user_id,
                    Task.due_date.is_not(None),
                    Task.due_date >= day_start,
                    Task.due_date <= day_end,
                )
            )
        )).scalars().all()
        actual_text, completion_rate = _summarize_tasks_for_day(tasks)
        entry.actual_text = actual_text
        entry.completion_rate = completion_rate
        if entry_day < datetime.utcnow().date():
            entry.locked = True

    await db.flush()
    refreshed = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    data = _template_payload(template, refreshed)
    return ApiResponse.success(data=data)


@router.put("/plan-template/export-mapping")
async def update_plan_template_export_mapping(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    body: PlanExportMappingPayload = None,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="当月计划模板不存在")
    payload = body.model_dump() if body else PlanExportMappingPayload().model_dump()
    template.export_mapping = json.dumps(payload, ensure_ascii=False)
    await db.flush()
    return ApiResponse.success(data=payload)


@router.get("/plan-template/validate")
async def validate_plan_template_before_export(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="当月计划模板不存在")
    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    missing_days = [e.day for e in entries if not (e.actual_text or "").strip() and not (e.manual_text or "").strip()]
    unlocked_past_days: list[int] = []
    today = datetime.utcnow().date()
    year, mon = map(int, month.split("-"))
    for e in entries:
        d = date(year, mon, e.day)
        if d < today and not e.locked:
            unlocked_past_days.append(e.day)
    return ApiResponse.success(data={
        "ok": len(missing_days) == 0 and len(unlocked_past_days) == 0,
        "missing_days": missing_days,
        "unlocked_past_days": unlocked_past_days,
    })


@router.get("/plan-template/export-submit-excel")
async def export_submit_plan_template_excel(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="当月计划模板不存在")
    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "上交版"
    mapping = {
        "date_col": "日期",
        "plan_col": "计划内容",
        "final_col": "最终上交内容",
        "rate_col": "完成率",
        "lock_col": "锁定",
    }
    if template.export_mapping:
        try:
            parsed = json.loads(template.export_mapping)
            if isinstance(parsed, dict):
                mapping.update({k: str(v) for k, v in parsed.items() if k in mapping})
        except Exception:
            pass
    has_slots = all(k in mapping for k in ("morning_col", "afternoon_col", "evening_col"))
    ws.append([f"{template.month} 备考计划上交表"])
    if has_slots:
        ws.append([
            mapping["date_col"],
            mapping["plan_col"],
            mapping["morning_col"],
            mapping["afternoon_col"],
            mapping["evening_col"],
            mapping["rate_col"],
            mapping["lock_col"],
        ])
    else:
        ws.append([mapping["date_col"], mapping["plan_col"], mapping["final_col"], mapping["rate_col"], mapping["lock_col"]])
    for e in entries:
        date_text = f"{template.month}-{str(e.day).zfill(2)}"
        final_text = "；".join([txt for txt in [e.actual_text or "", e.manual_text or ""] if txt]).strip("；")
        if has_slots:
            ws.append([
                date_text,
                e.planned_text or "",
                _extract_slot_text(e.actual_text, "上午"),
                _extract_slot_text(e.actual_text, "下午"),
                _extract_slot_text(e.actual_text, "晚间"),
                f"{e.completion_rate}%" if e.completion_rate is not None else "",
                "是" if e.locked else "否",
            ])
        else:
            ws.append([date_text, e.planned_text or "", final_text, f"{e.completion_rate}%" if e.completion_rate is not None else "", "是" if e.locked else "否"])
    col_count = 7 if has_slots else 5
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
    ws.freeze_panes = "A3"
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 34
    if has_slots:
        ws.column_dimensions["C"].width = 26
        ws.column_dimensions["D"].width = 26
        ws.column_dimensions["E"].width = 26
        ws.column_dimensions["F"].width = 10
        ws.column_dimensions["G"].width = 8
    else:
        ws.column_dimensions["C"].width = 44
        ws.column_dimensions["D"].width = 10
        ws.column_dimensions["E"].width = 8
    title_font = Font(name="Microsoft YaHei", bold=True, size=14)
    header_font = Font(name="Microsoft YaHei", bold=True, size=10)
    body_font = Font(name="Microsoft YaHei", size=10)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_wrap = Alignment(horizontal="left", vertical="top", wrap_text=True)
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill(fill_type="solid", start_color="EEF2FF", end_color="EEF2FF")
    ws["A1"].font = title_font
    ws["A1"].alignment = center
    for col in range(1, col_count + 1):
        hc = ws.cell(row=2, column=col)
        hc.font = header_font
        hc.alignment = center
        hc.fill = header_fill
        hc.border = border
    for row in range(3, ws.max_row + 1):
        for col in range(1, col_count + 1):
            c = ws.cell(row=row, column=col)
            c.font = body_font
            c.border = border
            if has_slots:
                c.alignment = center if col in (1, 6, 7) else left_wrap
            else:
                c.alignment = center if col in (1, 4, 5) else left_wrap
        ws.row_dimensions[row].height = 42
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    xlsx_bytes = output.read()
    filename = f"{template.name}_{template.month}_submit.xlsx".replace(" ", "_")
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.put("/plan-template/entry/{entry_id}")
async def update_plan_template_entry(
    entry_id: str,
    body: PlanEntryUpdatePayload,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskPlanEntry, TaskPlanTemplate)
        .join(TaskPlanTemplate, TaskPlanEntry.template_id == TaskPlanTemplate.id)
        .where(and_(TaskPlanEntry.id == entry_id, TaskPlanTemplate.user_id == user_id))
    )
    row = result.first()
    if not row:
        raise HTTPException(status_code=404, detail="计划条目不存在")
    entry: TaskPlanEntry = row[0]
    if body.planned_text is not None:
        entry.planned_text = body.planned_text
    if body.manual_text is not None:
        entry.manual_text = body.manual_text
    await db.flush()
    return ApiResponse.success(data={"id": entry.id, "planned_text": entry.planned_text, "manual_text": entry.manual_text})


@router.get("/plan-template/export-excel")
async def export_plan_template_excel(
    month: str = Query(..., pattern=r"^\d{4}-\d{2}$"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(TaskPlanTemplate).where(
            and_(TaskPlanTemplate.user_id == user_id, TaskPlanTemplate.month == month)
        )
    )
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="当月计划模板不存在")
    entries = (await db.execute(
        select(TaskPlanEntry).where(TaskPlanEntry.template_id == template.id).order_by(TaskPlanEntry.day.asc())
    )).scalars().all()
    xlsx_bytes = _build_plan_excel(template, entries)
    filename = f"{template.name}_{template.month}.xlsx".replace(" ", "_")
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{task_id}")
async def get_task(
    task_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Task).where(and_(Task.id == task_id, Task.user_id == user_id)))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return ApiResponse.success(data=TaskOut.model_validate(task).model_dump())


@router.put("/{task_id}")
async def update_task(
    task_id: str,
    body: TaskUpdate,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Task).where(and_(Task.id == task_id, Task.user_id == user_id)))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    update_data = body.model_dump(exclude_unset=True)
    if "tags" in update_data and update_data["tags"] is not None:
        update_data["tags"] = json.dumps(update_data["tags"])
    if "subtasks" in update_data and update_data["subtasks"] is not None:
        update_data["subtasks"] = json.dumps([s if isinstance(s, dict) else s.model_dump() for s in update_data["subtasks"]])

    if update_data.get("status") == "completed" and task.status != "completed":
        update_data["completed_at"] = datetime.utcnow()

    for key, value in update_data.items():
        setattr(task, key, value)

    await db.flush()
    await db.refresh(task)
    return ApiResponse.success(data=TaskOut.model_validate(task).model_dump())


@router.delete("/{task_id}")
async def delete_task(
    task_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Task).where(and_(Task.id == task_id, Task.user_id == user_id)))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    await db.delete(task)
    return ApiResponse.success(message="已删除")
